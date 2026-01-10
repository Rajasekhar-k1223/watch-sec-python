import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def md_to_docx(md_path, docx_path):
    print(f"Converting {md_path} -> {docx_path}")
    document = Document()
    
    # Simple Markdown Parser
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    table_mode = False
    table_data = []
    
    for line in lines:
        line = line.strip()
        
        # 1. Headings
        if line.startswith('#'):
            level = len(line.split()[0])
            text = line.lstrip('#').strip()
            document.add_heading(text, level=min(level, 9))
            
        # 2. Horizontal Rule
        elif line.startswith('---'):
            document.add_paragraph('__________________________________________________', style='Normal')
            
        # 3. Tables (Basic Detection)
        elif line.startswith('|'):
            if not table_mode:
                table_mode = True
                table_data = [] # Reset
            
            # Skip separator line | :--- |
            if '---' in line:
                continue
                
            cells = [c.strip() for c in line.split('|') if c]
            table_data.append(cells)
            
        # 4. Normal Text (Handle Table Exit)
        else:
            if table_mode and not line.startswith('|'):
                # Render Table
                if table_data:
                    cols = len(table_data[0])
                    table = document.add_table(rows=len(table_data), cols=cols)
                    table.style = 'Table Grid'
                    for i, row in enumerate(table_data):
                        # Ensure row has correct number of cells
                        match_cols = min(len(row), cols)
                        for j in range(match_cols):
                            table.cell(i, j).text = row[j]
                table_mode = False
                
            if line:
                p = document.add_paragraph()
                
                # Bold simple **text**
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)

    document.save(docx_path)

# Paths
base_dir = "/root/.gemini/antigravity/brain/80f562b8-1dc5-4c67-85bb-89353b4f5b8e"
files = [
    "SRS.md", 
    "Product_Marketing.md", 
    "Demo_Presentation.md", 
    "Demo_Video_Script.md"
]

for file in files:
    md_path = os.path.join(base_dir, file)
    docx_path = md_path.replace(".md", ".docx")
    try:
        if os.path.exists(md_path):
            md_to_docx(md_path, docx_path)
            print(f"Success: {docx_path}")
        else:
            print(f"Skipping missing: {md_path}")
    except Exception as e:
        print(f"Error converting {file}: {e}")
